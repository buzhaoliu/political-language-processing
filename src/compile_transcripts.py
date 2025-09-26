#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
compile_transcripts.py

Purpose
-------
Scan DOCX interviews in:
  /Users/buzhaoliu/Developer/NLP/Reviewed Transcription
Extract metadata (Corporator Name, Date, Location) and normalize the dialogue
into alternating Q_n / R_n columns, then write a single Excel file.

Output format (two rows per interview):
  row i   : headers [Corporator Name, Date, Location, Q_1, R_1, Q_2, R_2, ...]
  row i+1 : values  [name, date, location, q1_text, r1_text, q2_text, r2_text, ...]

Notes
-----
- Robust speaker detection across many variants (Interviewer/I:, Corporator/R:, etc.).
- Falls back to filename tokens when metadata is missing.
- If an interview starts with the corporator, Q_1 is blank and R_1 holds the first reply.
- Highlights empty metadata cells (gray) for quick manual review.

Usage
-----
python compile_transcripts.py \
  --docs "/Users/buzhaoliu/Developer/NLP/Reviewed Transcription" \
  --out  "/Users/buzhaoliu/Developer/NLP/interview_data_raw.xlsx"
"""

from __future__ import annotations
import argparse
import os
import re
from pathlib import Path
from typing import List, Tuple, Dict

import docx
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

# ---------- Defaults (your current layout) ----------
DEFAULT_DOCS = "/Users/buzhaoliu/Developer/NLP/Reviewed Transcription"
DEFAULT_OUT  = "/Users/buzhaoliu/Developer/NLP/interview_data_raw.xlsx"

# Variants seen in transcripts
INTERVIEWER_KEYS = [
    "[Interviewer]:", "Interviewer:", "I :", "I:", "I1 :", "I1:", "[Interviewer]", "I - "
]
CORPORATOR_KEYS = [
    "[Corporator]:", "Corporator:", "R :", "R:", "Ex-Corporator:", "[Corporator]", "R - "
]

META_FIELDS = ("Corporator Name", "Date", "Location")

# ---------- Helpers ----------
def get_docx_files(root: Path) -> List[Path]:
    return [p for p in root.rglob("*.docx") if not p.name.startswith("~$")]

def parse_metadata(doc: docx.Document, filename: str) -> Dict[str, str]:
    """
    Extract metadata from 'Key: Value' lines; if missing, fall back to filename
    tokens like 'Name_Date_Location' (best-effort).
    """
    md = {k: "" for k in META_FIELDS}
    for para in doc.paragraphs:
        t = para.text.strip()
        for key in META_FIELDS:
            if t.lower().startswith(key.lower() + ":"):
                md[key] = t.split(":", 1)[1].strip()

    # Fallback: filename like 'Name_2015_Lucknow' -> Name/2015/Lucknow
    if any(v == "" for v in md.values()):
        base = Path(filename).stem
        parts = [x.strip() for x in re.split(r"[_\-\s]+", base) if x.strip()]
        if parts:
            if not md["Corporator Name"]:
                md["Corporator Name"] = parts[0]
            if len(parts) > 1 and not md["Date"]:
                md["Date"] = parts[1]
            if len(parts) > 2 and not md["Location"]:
                md["Location"] = " ".join(parts[2:])
    return md

def split_speaker_segments(doc: docx.Document) -> List[str]:
    """
    Combine paragraphs, then split into segments at any speaker token
    to make downstream Q/R extraction more resilient.
    """
    full_text = " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    tokens = full_text.split()
    keys = INTERVIEWER_KEYS + CORPORATOR_KEYS

    segments, cur = [], ""
    for w in tokens:
        if any(w.startswith(k) for k in keys):
            if cur.strip():
                segments.append(cur.strip())
            cur = w + " "
        else:
            cur += w + " "
    if cur.strip():
        segments.append(cur.strip())
    return segments

def parse_qr_from_segments(segments: List[str]) -> Tuple[List[str], List[str], bool]:
    """Return (questions, answers, starts_with_corporator)."""
    q, r = [], []
    starts_with_corp = False
    if segments:
        starts_with_corp = any(segments[0].startswith(k) for k in CORPORATOR_KEYS)

    for seg in segments:
        if any(seg.startswith(k) for k in INTERVIEWER_KEYS):
            for kw in INTERVIEWER_KEYS:
                if seg.startswith(kw):
                    q.append(seg[len(kw):].strip())
                    break
        elif any(seg.startswith(k) for k in CORPORATOR_KEYS):
            for kw in CORPORATOR_KEYS:
                if seg.startswith(kw):
                    r.append(seg[len(kw):].strip())
                    break
    return q, r, starts_with_corp

def compile_one_file(fp: Path) -> List[List[str]]:
    """Produce two rows (header, values) for a single interview DOCX."""
    doc = docx.Document(str(fp))
    meta = parse_metadata(doc, fp.name)
    segments = split_speaker_segments(doc)
    questions, answers, starts_with_corp = parse_qr_from_segments(segments)

    max_pairs = max(len(questions), len(answers))
    headers = ["", "", ""] + [f"{'Q' if i % 2 == 0 else 'R'}_{(i//2)+1}" for i in range(2 * max_pairs)]
    values  = [meta["Corporator Name"], meta["Date"], meta["Location"]]

    q_i, r_i = 0, 0
    # If conversation starts with corporator, first Q is blank, R_1 holds the first reply
    if starts_with_corp:
        values += ["", answers[0] if answers else ""]
        r_i = 1

    while q_i < len(questions) or r_i < len(answers):
        q_val = questions[q_i] if q_i < len(questions) else ""
        r_val = answers[r_i]   if r_i < len(answers)   else ""
        if not (starts_with_corp and q_i == 0 and r_i == 0):
            values += [q_val, r_val]
        q_i += 1
        r_i += 1

    headers[:3] = list(META_FIELDS)
    return [headers, values]

# ---------- Main ----------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docs", default=DEFAULT_DOCS, help="Folder containing interview DOCX files")
    ap.add_argument("--out",  default=DEFAULT_OUT,  help="Path to output Excel file")
    args = ap.parse_args()

    doc_root = Path(args.docs)
    files = get_docx_files(doc_root)
    if not files:
        print(f"⚠️ No .docx files found under: {doc_root}")
        return

    all_rows: List[List[str]] = []
    first = True
    for fp in sorted(files):
        rows = compile_one_file(fp)
        if first:
            all_rows.extend(rows)      # keep the header row the first time
            first = False
        else:
            # blank out meta headers on subsequent blocks to avoid repetition
            rows[0][:3] = ["", "", ""]
            all_rows.extend(rows)

    df = pd.DataFrame(all_rows)
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    df.to_excel(out_path, index=False, header=False)

    # Highlight empty metadata cells in value rows for quick QA
    wb = load_workbook(out_path)
    ws = wb.active
    gray = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
    for r in range(2, ws.max_row + 1, 2):  # value rows only (1-based)
        for c in range(1, 4):  # Corporator Name, Date, Location
            cell = ws.cell(row=r, column=c)
            if not (cell.value and str(cell.value).strip()):
                cell.fill = gray
    wb.save(out_path)

    print(f"✅ Compiled spreadsheet saved to: {out_path}")
    print(f"📁 Source folder: {doc_root}")
    print(f"🗒️ Interviews processed: {len(files)}")

if __name__ == "__main__":
    main()
